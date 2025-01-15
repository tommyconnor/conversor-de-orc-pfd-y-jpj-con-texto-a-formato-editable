import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Progressbar
import os
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image
import pytesseract
import pdfplumber

class OrtegaMiJefaApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Ortega Mi Jefa")
        self.doc_counter = 1
        self.load_doc_counter()
        
        self.create_widgets()

    def create_widgets(self):
        self.load_button = tk.Button(self.root, text="Cargar Documento", command=self.load_document)
        self.load_button.pack(pady=10)

        self.convert_button = tk.Button(self.root, text="Comenzar Conversión", command=self.start_conversion)
        self.convert_button.pack(pady=10)

        self.progress = Progressbar(self.root, orient=tk.HORIZONTAL, length=200, mode='determinate')
        self.progress.pack(pady=10)

    def load_document(self):
        # Permitir cargar archivos PDF y JPG
        self.filepath = filedialog.askopenfilename(
            initialdir=os.path.expanduser("~/Documents"),
            filetypes=[
                ("Archivos PDF", "*.pdf"),
                ("Imágenes JPG", "*.jpg"),
                ("Imágenes PNG", "*.png"),
                ("Todos los archivos", "*.*")
            ]
        )
        if self.filepath:
            messagebox.showinfo("Archivo Cargado", f"Archivo cargado: {self.filepath}")

    def start_conversion(self):
        if not hasattr(self, 'filepath'):
            messagebox.showerror("Error", "Primero cargue un documento.")
            return

        format_choice = messagebox.askquestion("Formato de Conversión", "¿Convertir a PDF Normal?\nSi selecciona No, se convertirá a Word Editable.")
        if format_choice == 'yes':
            self.convert_to_pdf()
        else:
            format_choice = messagebox.askquestion("Formato de Conversión", "¿Convertir a Excel Editable?\nSi selecciona No, se convertirá a Word Editable.")
            if format_choice == 'yes':
                self.convert_to_excel()
            else:
                self.convert_to_word()

    def convert_to_pdf(self):
        output_filepath = os.path.expanduser(f"~/Documents/proyecto_{self.doc_counter}.pdf")
        c = canvas.Canvas(output_filepath, pagesize=letter)

        if self.filepath.endswith('.pdf'):
            with pdfplumber.open(self.filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        c.drawString(100, 750, text.strip())  # Ajustar la posición según sea necesario

        c.save()
        self.update_progress(100)
        messagebox.showinfo("Conversión Completa", f"El documento ha sido convertido a PDF y guardado en {output_filepath}")
        self.doc_counter += 1
        self.save_doc_counter()

    def convert_to_word(self):
        output_filepath = os.path.expanduser(f"~/Documents/proyecto_{self.doc_counter}.docx")
        doc = Document()

        if self.filepath.endswith('.jpg') or self.filepath.endswith('.png'):
            # Convertir imagen a texto usando OCR
            text = pytesseract.image_to_string(Image.open(self.filepath))
            doc.add_paragraph(text)

        elif self.filepath.endswith('.pdf'):
            with pdfplumber.open(self.filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text.strip())

        doc.save(output_filepath)
        self.update_progress(100)
        messagebox.showinfo("Conversión Completa", f"El documento ha sido convertido a Word y guardado en {output_filepath}")
        self.doc_counter += 1
        self.save_doc_counter()

    def convert_to_excel(self):
        output_filepath = os.path.expanduser(f"~/Documents/proyecto_{self.doc_counter}.xlsx")
        
        data = []  # Aquí almacenarás los datos para Excel

        if self.filepath.endswith('.jpg') or self.filepath.endswith('.png'):
            text = pytesseract.image_to_string(Image.open(self.filepath))
            lines = text.splitlines()
            for line in lines:
                parts = line.strip().split(' ')
                data.append(parts)

        elif self.filepath.endswith('.pdf'):
            with pdfplumber.open(self.filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.splitlines()
                        for line in lines:
                            parts = line.strip().split(' ')
                            data.append(parts)

        # Estructurar datos para Excel según el formato solicitado
        structured_data = []
        
        for row in data:
            structured_row = []
            for i, value in enumerate(row):
                structured_row.append(value + f' {i + 1}')  # Añadir el índice como sufijo
            structured_data.append(structured_row)

        df = pd.DataFrame(structured_data)
        
        # Guardar en Excel
        df.to_excel(output_filepath, index=False, header=False)

        self.update_progress(100)
        messagebox.showinfo("Conversión Completa", f"El documento ha sido convertido a Excel y guardado en {output_filepath}")
        self.doc_counter += 1
        self.save_doc_counter()

    def update_progress(self, value):
        self.progress['value'] = value
        self.root.update_idletasks()

    def load_doc_counter(self):
        try:
            with open("doc_counter.txt", "r") as f:
                self.doc_counter = int(f.read().strip())
        except FileNotFoundError:
            self.doc_counter = 1

    def save_doc_counter(self):
        with open("doc_counter.txt", "w") as f:
            f.write(str(self.doc_counter))

if __name__ == "__main__":
    root = tk.Tk()
    app = OrtegaMiJefaApp(root)
    root.mainloop()  # Esta línea es crucial para mostrar la ventana.



